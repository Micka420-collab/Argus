#!/usr/bin/env python3
"""
Génère docs/Argus_Documentation.docx — documentation technique complète.
Usage : python scripts/gen_documentation.py
Dépendance : python-docx
"""
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x9E)
MUTED  = RGBColor(0x55, 0x5B, 0x66)
CODE_BG = "F2F4F7"
HEAD_BG = "1F4E9E"


# ---------------------------------------------------------------------------
# Helpers OOXML
# ---------------------------------------------------------------------------
def _shade(el, fill):
    pr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    pr.append(shd)


def code_block(doc, text):
    p = doc.add_paragraph()
    _shade(p._p, CODE_BG)
    pf = p.paragraph_format
    pf.left_indent = Pt(8); pf.space_before = Pt(4); pf.space_after = Pt(8)
    for i, line in enumerate(text.strip("\n").split("\n")):
        run = p.add_run(("" if i == 0 else "\n") + line)
        run.font.name = "Consolas"; run.font.size = Pt(9)
        # police mono côté east-asian aussi
        rpr = run._r.get_or_add_rPr(); rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rpr.append(rf)
        rf.set(qn("w:ascii"), "Consolas"); rf.set(qn("w:hAnsi"), "Consolas")
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.size = Pt(9.5)
        _shade(hdr[i]._tc, HEAD_BG)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    return t


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text, muted=False):
    par = doc.add_paragraph()
    run = par.add_run(text)
    if muted:
        run.font.color.rgb = MUTED
    return par


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_toc(doc):
    par = doc.add_paragraph()
    run = par.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Sommaire — clic droit → « Mettre à jour les champs » (ou F9)."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for e in (f1, instr, f2, t, f3):
        run._r.append(e)


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------
doc = Document()

# Style de base + couleurs de titres
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)
for i in range(1, 4):
    try:
        doc.styles[f"Heading {i}"].font.color.rgb = ACCENT
    except KeyError:
        pass

# ---- Page de titre ----
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(120)
r = t.add_run("ARGUS"); r.bold = True; r.font.size = Pt(54); r.font.color.rgb = ACCENT
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("SOC autonome & post-quantique"); r.font.size = Pt(18); r.font.color.rgb = MUTED
tg = doc.add_paragraph(); tg.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tg.add_run("« Le SOC qui s'investigue lui-même »"); r.italic = True; r.font.size = Pt(13)
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(60)
r = meta.add_run(f"Documentation technique — v3.0.0\n{date.today().isoformat()}\n"
                 "Licence PolyForm Noncommercial 1.0.0 (source-available, non-commercial)")
r.font.size = Pt(11); r.font.color.rgb = MUTED
doc.add_page_break()

# ---- Sommaire ----
h(doc, "Sommaire", 1)
add_toc(doc)
doc.add_page_break()

# ============================================================ 1
h(doc, "1. Présentation & vision", 1)
p(doc, "Argus est une plateforme SOC (Security Operations Center) open-source, "
       "auto-hébergeable et résistante au quantique. Elle s'inspire de quatre acteurs du "
       "marché et recrée leur valeur en open-source au sein d'une seule plateforme.")
table(doc, ["Pilier", "Inspiration", "Capacité"], [
    ["A", "Qevlar AI", "Analyste IA autonome — investigation de bout en bout, verdict, rapport"],
    ["B", "CryptoNext", "Posture post-quantique — TLS hybride ML-KEM, JWT hybride ML-DSA, CBOM"],
    ["C", "Snowpack", "Réseau anonymisé — égress OSINT via Tor/proxy, passerelle anon-gateway"],
    ["D", "YesWeHack", "Bug-bounty / VDP — soumission, triage, CVSS, récompenses"],
])
p(doc, "Principe directeur (cœur de l'approche Qevlar) : le verdict est calculé de façon "
       "déterministe (Python) ; le LLM est borné à la rédaction du rapport — il ne décide jamais.")

# ============================================================ 2
h(doc, "2. Les quatre piliers", 1)

h(doc, "2.1 Pilier A — Analyste IA autonome", 2)
p(doc, "Graphe d'investigation déterministe (style LangGraph, sans dépendance lourde). "
       "Chaque nœud enrichit un état partagé et journalise une trace auditable :")
code_block(doc, "ingest → enrich(OSINT) → correlate(OpenSearch, beaconing) → retrieve_feedback(RAG)\n"
                "      → score(déterministe) → decide → report(LLM borné)\n"
                "      → propose_actions(human-gated) → persist → escalate")
bullets(doc, [
    "Verdict 3 états (malicious / benign / inconclusive) + confiance — services/scoring.py.",
    "LLM borné (Ollama / Claude / repli heuristique) — services/llm.py.",
    "RAG : les corrections analyste sont réinjectées dans les cas similaires — services/feedback.py.",
    "Actions destructrices jamais automatiques (human-in-the-loop).",
])

h(doc, "2.2 Pilier B — Posture post-quantique", 2)
bullets(doc, [
    "Transport : TLS hybride X25519+ML-KEM-768 (FIPS 203) au niveau Nginx ; repli X25519.",
    "Authentification : JWT hybride Ed25519 + ML-DSA-65 (FIPS 204), activé par PQC_JWT ; repli HS256.",
    "Gouvernance : scanner de readiness + CBOM classant chaque algorithme (sûr / hybride / vulnérable).",
])

h(doc, "2.3 Pilier C — Réseau anonymisé", 2)
bullets(doc, [
    "Client HTTP sortant proxy-aware : OSINT routé par Tor/proxy quand OSINT_ANON=true.",
    "Passerelle anon-gateway : Tor SOCKS (:9050) + rotation d'identité NEWNYM (:9052).",
    "Scaffolding mesh WireGuard + Rosenpass (ML-KEM) pour un plan de management sans port entrant.",
])

h(doc, "2.4 Pilier D — VDP / Bug-Bounty", 2)
bullets(doc, [
    "Soumission de rapports, machine à états, triage assisté (dédup + résumé IA).",
    "Calculateur CVSS v3.1 maison + grille de récompenses (sévérité × valeur métier).",
    "Rôles dédiés : researcher, triager.",
])

h(doc, "2.5 Extensions — ASM/CTEM & Intégrations", 2)
bullets(doc, [
    "ASM / CTEM : registre d'exposition priorisé CVSS × EPSS × KEV × valeur métier "
    "(enrichissement FIRST EPSS + CISA KEV ; plancher critique pour les CVE KEV).",
    "Webhooks sortants signés HMAC (n8n/Slack/Jira) sur verdict IA, VDP accepté, finding KEV.",
])

# ============================================================ 3
h(doc, "3. Architecture", 1)
p(doc, "Point d'entrée unique : Nginx sur 443. Tous les services internes restent sur le "
       "réseau Docker privé soc-net (172.20.0.0/24) et ne sont jamais exposés sur l'hôte.")
table(doc, ["Service", "Techno", "Rôle"], [
    ["nginx", "nginx 1.29-alpine (OpenSSL ≥ 3.5)", "Reverse proxy TLS hybride PQC"],
    ["soc-frontend", "React 18 + Tailwind", "UI « Obsidian Signal » + landing"],
    ["soc-api", "FastAPI (Python 3.11)", "REST, WebSocket, moteur d'alerting, agent IA"],
    ["opensearch", "OpenSearch 2.11", "Stockage & recherche"],
    ["wazuh-manager", "Wazuh 4.7", "SIEM/EDR, Active Response"],
    ["suricata", "Suricata", "IDS/IPS réseau (host)"],
    ["redis", "redis 7", "Cache, déduplication, sessions"],
    ["ollama", "ollama (profil ai)", "LLM local pour l'analyste IA"],
    ["anon-gateway", "Tor + WireGuard (profil anon)", "Égress anonymisé / mesh PQC"],
    ["n8n", "n8n (profil n8n)", "Orchestration / sync tickets"],
])
h(doc, "3.1 Stockage", 2)
table(doc, ["Données", "Index / clé"], [
    ["Alertes Wazuh", "wazuh-alerts-*"],
    ["Événements Suricata", "soc-suricata-*"],
    ["Investigations IA", "soc-investigations"],
    ["Feedback analyste (RAG)", "soc-ai-feedback"],
    ["Rapports VDP", "soc-vdp-reports (+ -programs)"],
    ["Exposition ASM/CTEM", "soc-exposure-assets / -findings"],
    ["Clés ML-DSA (JWT PQC)", "volume pqc_keys"],
])

# ============================================================ 4
h(doc, "4. Installation", 1)
h(doc, "4.1 Installeur guidé (recommandé)", 2)
p(doc, "L'installeur vérifie Docker, génère tous les secrets, crée le certificat TLS, "
       "configure les options selon vos réponses, démarre la stack et affiche l'URL + le mot de passe admin.")
p(doc, "Linux / macOS :")
code_block(doc, "curl -fsSL https://raw.githubusercontent.com/Micka420-collab/Argus/main/install.sh | bash")
p(doc, "Windows (Docker Desktop) :")
code_block(doc, "git clone https://github.com/Micka420-collab/Argus.git ; cd Argus\n"
                "powershell -ExecutionPolicy Bypass -File .\\install.ps1")
p(doc, "Sans question (valeurs par défaut) : ARGUS_YES=1 ./install.sh", muted=True)
h(doc, "4.2 Installation manuelle", 2)
code_block(doc, "git clone https://github.com/Micka420-collab/Argus.git\n"
                "cd Argus\ncp .env.example .env   # éditer les secrets\ndocker compose up -d")
p(doc, "Console : https://localhost · Présentation : https://localhost/welcome")
p(doc, "Prérequis : Docker + Compose v2 ; vm.max_map_count=262144 (OpenSearch).", muted=True)

# ============================================================ 5
h(doc, "5. Maintenance — cycle de vie", 1)
table(doc, ["Script", "Rôle"], [
    ["install.sh / .ps1", "Installation guidée (secrets, TLS, options, démarrage)"],
    ["update.sh / .ps1", "git pull + rebuild + restart (préserve données & secrets)"],
    ["backup.sh / .ps1", "Sauvegarde volumes + .env + certificats → backups/argus-<date>/"],
    ["restore.sh / .ps1", "Restaure une sauvegarde (écrase les données)"],
    ["uninstall.sh / .ps1", "Arrête la stack ; suppression optionnelle données/images"],
])
p(doc, "Sauvegarde automatisable via cron :")
code_block(doc, "0 3 * * * cd /opt/Argus && ARGUS_YES=1 ./backup.sh")

# ============================================================ 6
h(doc, "6. Configuration (variables d'environnement)", 1)
table(doc, ["Variable", "Défaut", "Description"], [
    ["SECRET_KEY", "(aléatoire)", "Obligatoire & persistant. Signature HS256 + graine Ed25519."],
    ["ADMIN_USERNAME / ADMIN_PASSWORD", "admin / …", "Compte admin créé au 1er démarrage."],
    ["OPENSEARCH_PASSWORD / REDIS_PASSWORD / WAZUH_API_PASSWORD", "—", "Mots de passe des services."],
    ["LLM_PROVIDER / LLM_MODEL", "none", "none | ollama | claude ; ex. qwen2.5:7b."],
    ["ANTHROPIC_API_KEY", "—", "Requis si LLM_PROVIDER=claude."],
    ["AI_AUTO_INVESTIGATE", "false", "Agent autonome sur alertes critiques."],
    ["PQC_JWT / PQC_KEYS_DIR", "false", "JWT hybride Ed25519(+ML-DSA) ; persistance des clés."],
    ["TLS_GROUPS", "X25519MLKEM768:X25519:secp256r1", "Groupes d'échange TLS (edge)."],
    ["OSINT_ANON / OUTBOUND_PROXY", "false / —", "Égress OSINT via Tor/proxy."],
    ["WEBHOOK_URL / WEBHOOK_SECRET", "—", "Sync tickets (n8n/Slack/Jira), events signés HMAC."],
    ["ABUSEIPDB_KEY / VIRUSTOTAL_KEY", "—", "Clés OSINT (comptes gratuits)."],
])

# ============================================================ 7
h(doc, "7. API REST", 1)
p(doc, "Base /api/v1. Authentification par JWT (en-tête Authorization: Bearer, ou cookie httpOnly). "
       "Swagger sur /docs quand DEBUG=true.")
h(doc, "7.1 Analyste IA", 2)
table(doc, ["Méthode & endpoint", "Description"], [
    ["GET /ai/reports", "Liste des investigations autonomes"],
    ["GET /ai/report/{id}", "Détail (verdict, preuves, trace)"],
    ["POST /ai/investigate/{alert_id}", "Lance l'agent sur une alerte"],
    ["POST /ai/investigate-ip/{ip}", "Lance l'agent sur une IP"],
    ["POST /ai/feedback/{id}", "Corriger/valider un verdict → RAG"],
])
h(doc, "7.2 Post-quantique & VDP & Exposition", 2)
table(doc, ["Méthode & endpoint", "Description"], [
    ["GET /crypto/readiness", "Posture post-quantique de la plateforme"],
    ["GET /crypto/inventory", "CBOM (handshakes TLS observés)"],
    ["GET /crypto/jwt-info", "Algorithme de signature JWT actif"],
    ["POST /vdp/reports", "Soumettre un rapport (CVSS auto)"],
    ["PATCH /vdp/reports/{id}/status", "Transition d'état (triage)"],
    ["GET /vdp/cvss?vector=", "Calcul CVSS v3.1"],
    ["GET /exposure/findings", "Findings priorisés (EPSS/KEV)"],
    ["POST /exposure/findings", "Ajouter un finding (enrichi)"],
])

# ============================================================ 8
h(doc, "8. Sécurité", 1)
h(doc, "8.1 Rôles (RBAC)", 2)
table(doc, ["Rôle", "Permissions"], [
    ["admin", "read, write, delete, admin, triage"],
    ["analyst", "read, write, triage"],
    ["triager", "read, triage"],
    ["researcher", "read, submit"],
    ["viewer", "read"],
])
h(doc, "8.2 JWT — classique vs post-quantique", 2)
bullets(doc, [
    "PQC_JWT=false (défaut) : HS256 (HMAC symétrique, quantum-résistant).",
    "PQC_JWT=true : Ed25519 (dérivé déterministe de SECRET_KEY) + ML-DSA-65 si liboqs présent.",
    "Vérification : Ed25519 obligatoire + ML-DSA si présent ; clé publique via /crypto/jwt-info.",
])
h(doc, "8.3 Garde-fous & durcissement", 2)
bullets(doc, [
    "Le verdict est déterministe ; le LLM ne décide jamais.",
    "Actions destructrices jamais auto-exécutées (requires_confirmation).",
    "Un seul port exposé (443) ; services internes isolés sur soc-net.",
    "En production : SECRET_KEY fort/persistant, DEBUG=false, certificat TLS valide, "
    "CORS restreint, UFW sur 1514/1515, sauvegardes chiffrées.",
])

# ============================================================ 9
h(doc, "9. Intégration continue (CI)", 1)
p(doc, "GitHub Actions (.github/workflows/ci.yml) sur push/PR vers main :")
bullets(doc, [
    "backend : pip install + compileall + smoke import du graphe complet + ruff (non bloquant).",
    "frontend : npm ci + build de production.",
    "scripts : ShellCheck + validation docker compose config.",
    "images : build des images api/frontend/nginx/anon-gateway (push sur main).",
])

# ============================================================ 10
h(doc, "10. Feuille de route", 1)
p(doc, "Livré : analyste IA autonome (graphe + RAG), post-quantique (TLS + JWT hybrides + CBOM), "
       "égress anonymisé + anon-gateway, VDP/Bug-Bounty, ASM/CTEM, webhooks/sync tickets, "
       "refonte UI + landing, CI, cycle de vie (install/update/backup/restore/uninstall).")
p(doc, "Phase 2/3 : déploiement du mesh WireGuard+Rosenpass (clés par site), scan actif ASM "
       "(nuclei/httpx), PKI mTLS interne.")

# ============================================================ 11
h(doc, "11. Licence", 1)
p(doc, "PolyForm Noncommercial License 1.0.0 — code source-available, libre pour tout usage "
       "non-commercial (perso, éducatif, recherche). Revente et usage commercial interdits sans "
       "licence commerciale séparée. Contact : micka.delcato.rp@gmail.com.")

doc.save("docs/Argus_Documentation.docx")
print("OK -> docs/Argus_Documentation.docx")
